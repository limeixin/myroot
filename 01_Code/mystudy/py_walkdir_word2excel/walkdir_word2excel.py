#! phthon
#encoding: utf-8


import os
import sys
import docx
import openpyxl
import re
import time


g_sheet = None
g_row = 1


def create_excel_title():
    g_sheet['A1'] = '简历名称'
    g_sheet['B1'] = '姓名'
    g_sheet['C1'] = '性别'
    g_sheet['D1'] = '年龄'
    g_sheet['E1'] = '工作经验'
    g_sheet['F1'] = '学历'
    g_sheet['G1'] = '联系方式'
    g_sheet['H1'] = '毕业学校'
    g_sheet['I1'] = '公司'
    g_sheet['J1'] = 'linux经验'
    g_sheet['K1'] = 'C语言经验'


def mark_filename(str):
    g_sheet.cell(row=g_row, column=1).value = str
def mark_name(str):
    g_sheet.cell(row=g_row, column=2).value = str
def mark_sex(str):
    g_sheet.cell(row=g_row, column=3).value = str
def mark_age(str):
    g_sheet.cell(row=g_row, column=4).value = str
def mark_years(str):
    g_sheet.cell(row=g_row, column=5).value = str
def mark_education(str):
    g_sheet.cell(row=g_row, column=6).value = str
def mark_telephone(str):
    g_sheet.cell(row=g_row, column=7).value = str
def mark_university(str):
    g_sheet.cell(row=g_row, column=8).value = str
def mark_companys(str):
    g_sheet.cell(row=g_row, column=9).value = str
def mark_linux(str):
    g_sheet.cell(row=g_row, column=10).value = str
def mark_C_language(str):
    g_sheet.cell(row=g_row, column=11).value = str
def mark_other(str):
    g_sheet.cell(row=g_row, column=12).value = str


def zhilian(filename):
    global g_row
    g_row += 1

    doc = docx.Document(filename)
    # i = 0
    # for para in doc.paragraphs:
        # re.search(, para.text)
        
        # print("paragraphs[%d] ====>" % i)
        # print(para.text)
        # i += 1

    # i = 0
    # for tab in doc.tables:
        # print("tables[%d]" % i)
        # for row in range(tab.rows.__len__()):
            # print("rows[%d]" % row)
            # for j in range(tab.rows[row].cells.__len__()):
                # print(tab.rows[row].cells[j].text)
        # i += 1

    #write to excel
    mark_filename(filename)

    myre = re.search(r'智联招聘_([\u4e00-\u9fa5]{2,4})_中文', filename)
    if myre:
        mark_name(myre.group(1))

    sexre = re.search(r'男', doc.tables[1].rows[1].cells[0].text)
    if sexre:
        mark_sex('男')
    else:
        mark_sex('女')

    agere = re.search(r'(\d{2})岁', doc.tables[1].rows[1].cells[0].text)
    if agere:
        mark_age(agere.group(1))

    years = re.search(r'(\d{1,2})年工作经验', doc.tables[1].rows[1].cells[0].text)
    if years:
        mark_years(years.group(1))

    edu1 = re.search(r'本科', doc.tables[1].rows[1].cells[0].text)
    if edu1:
        mark_education('本科')
    else:
        edu2 = re.search(r'硕士', doc.tables[1].rows[1].cells[0].text)
        if edu2:
            mark_education('硕士')

    tel = re.search(r'\d{11}', doc.tables[1].rows[2].cells[0].text)
    if tel:
        mark_telephone(tel.group())

    i = 3
    universitys = ''
    companys = ''
    havelinux = 0
    havec = 0
    for tab in doc.tables:
        # print("tables[%d]" % i)
        for row in range(tab.rows.__len__()):
            # print("rows[%d]" % row)
            for j in range(tab.rows[row].cells.__len__()):
                #print(tab.rows[row].cells[j].text)
                university = re.search(r'([\u4e00-\u9fa5]{2,15}大学)', tab.rows[row].cells[j].text)
                if university:
                    universitys += university.group(1) + '\n'

                company = re.search(r'([\u4e00-\u9fa5]{2,20}公司)', tab.rows[row].cells[j].text)
                if company:
                    companys += company.group(1) + '\n'

                linux = re.search(r'linux', tab.rows[row].cells[j].text, re.I)
                if linux:
                    havelinux = 1

                clang = re.search(r'C语言', tab.rows[row].cells[j].text)
                if clang:
                    havec = 1
        i += 1
    if universitys != '':
        mark_university(universitys)
    if companys != '':
        mark_companys(companys)

    # search all paragraphs
    for para in doc.paragraphs:
        if havelinux == 0:
            linux = re.search('linux', para.text, re.I)
            if linux:
                havelinux = 1    
        if havec == 0:
            clang = re.search(r'C语言', tab.rows[row].cells[j].text)
            if clang:
                havec = 1

    if havelinux == 1:
        mark_linux('是')
    else:
        mark_linux('否')
    if havec == 1:
        mark_C_language('是')
    else:
        mark_C_language('否')


def job51(filename):
    pass
    # global g_row
    # g_row += 1
    # doc = docx.Document(filename)

    # i = 0
    # for tab in doc.tables:
        # print("tables[%d]" % i)
        # i += 1
        # for row in range(tab.rows.__len__()):
            # print("rows[%d]" % row)
            # for j in range(tab.rows[row].cells.__len__()):
                # print("tab.rows[%d].cells[%d]" % (row,j))
                # print(tab.rows[row].cells[j].tables)

                # secondlen = len(tab.rows[row].cells[j].tables)
                # print("tab.rows[%d].cells[%d] --> tables len %d" % (row,j, secondlen))

                # if secondlen > 0:
                    # for x in range(secondlen):
                        # sts = tab.rows[row].cells[j].tables[x]
                        # for y in range(sts.rows.__len__()):
                            # print("sts[%d]" % y)
                            # for z in range(sts.rows[y].cells.__len__()):
                                # print(sts.rows[y].cells[z].text)
                # else:
                    # print(tab.rows[row].cells[j].text
    # for table in doc.tables:
        # for row in table.rows:
            # for cell in row.cells:
                # print(cell.text)

    # print("-----------------------------------------------------------")
    # for para in doc.paragraphs:
        # print("paragraphs[] ====>" )
        # print(para.text)
        # print("===========")
                        



def liepin(filename):
    global g_row
    g_row += 1
    doc = docx.Document(filename)
    mark_filename(filename)
    mark_name(doc.tables[2].rows[0].cells[1].text)
    mark_sex(doc.tables[2].rows[0].cells[3].text)
    mark_telephone(doc.tables[2].rows[1].cells[1].text)
    mark_age(doc.tables[2].rows[1].cells[3].text)
    mark_years(doc.tables[2].rows[3].cells[1].text)
    mark_education(doc.tables[2].rows[2].cells[3].text)

    universitys = ''
    companys = ''
    havelinux = 0
    havec = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                #print(cell.text)
                university = re.search(r'([\u4e00-\u9fa5]{2,15}大学)', cell.text)
                if university:
                    universitys += university.group(1) + '\n'

                company = re.search(r'([\u4e00-\u9fa5]{2,20}公司)', cell.text)
                if company:
                    companys += company.group(1) + '\n'

                if ('linux' in cell.text) or ('Linux' in cell.text):
                    havelinux = 1

                #clang = re.search(r'C语言', cell.text)
                if r'C语言' in cell.text:
                    havec = 1

    if universitys != '':
        mark_university(universitys)
    if companys != '':
        mark_companys(companys)
    if havelinux == 1:
        mark_linux('是')
    else:
        mark_linux('否')
    if havec == 1:
        mark_C_language('是')
    else:
        mark_C_language('否')


def others(filename):
    print("others")




def resume_summary(filename):
    if not filename.endswith("docx"):
        print("%s is not a docx file." % filename)
        return

    if '猎聘' in os.path.basename(filename):
        liepin(filename)
    elif '51job' in os.path.basename(filename):
        job51(filename)
    elif '智联' in os.path.basename(filename):
        zhilian(filename)
    else:
        others(filename)
        


def walk_file(pathname):
    filelist = os.listdir(pathname)
    for filename in filelist:
        fullname = os.path.join(pathname, filename)
        if os.path.isdir(fullname):
            walk_file(fullname)
        elif os.path.isfile(fullname):
            #print(fullname)
            # job of specific
            resume_summary(fullname)
        else:
            print("%s is not normal file")


def main():
    if sys.version_info[0] != 3 :
        print("This program use python3, but your environment is python%d" % sys.version_info[0])
        sys.exit()

    # check input arguments for file path
    if len(sys.argv) != 2 :
        print("Please input a path of directory.")
        sys.exit()
    if not os.path.isdir(sys.argv[1]):
        print("You input %s is not a correct directory." % sys.argv[1])
        sys.exit()

    fullpath = sys.argv[1]

    # delete delimiter of tail
    while fullpath.endswith('\\') or fullpath.endswith('/'):
        fullpath = fullpath[:-1]

    # create a new excel file.
    global g_sheet
    wb = openpyxl.Workbook() #here not check, need if wb == None check.
    #g_sheet = wb.get_active_sheet()
    g_sheet = wb.active
    g_sheet.title = '简历摘要'

    # excel title line
    create_excel_title()

    # process job
    walk_file(fullpath)

    # save excel file
    xlsx = os.path.join(fullpath, '简历提取摘要' + time.strftime('%Y%m%d-%H%M%S') + '.xlsx')
    wb.save(xlsx)


if __name__ == "__main__":
    main()